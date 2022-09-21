'''
1.In what modes should the PdfFileReader() and PdfFileWriter() File objects will be opened?
Ans.These files will be opened in binary mode.
    read binary (rb) for PdfFileREader()
    write binary (wb) PdfFileWriter()

2.From a PdfFileReader object, how do you get a Page object for page 5?
Ans.Calling getPage(4)

import PyPDF2 as pdf
pdfFileObj = open("cardealer.pdf",'rb')
pdfReader = pdf.PdfFileReader(pdfFileObj)
pageObj = pdfReader.getPage(4)
pageObj.extractText()

3.What PdfFileReader variable stores the number of pages in the PDF document?
Ans.   .numPages

4.If a PdfFileReader object’s PDF is encrypted with the password swordfish, what must you do
before you can obtain Page objects from it?
Ans.The pdf has to be decrypted by calling .decrypt('swordfish')

5.What methods do you use to rotate a page?
Ans. pageObj.rotateClockwise(90)

6.What is the difference between a Run object and a Paragraph object?
Ans.
Run Objects : Runs are contiguous groups of characters within a paragraph with the same style
Paragraph Object : A document contains multiple paragraphs.
A paragraph begins on a new line and contains multiple runs.

7.How do you obtain a list of Paragraph objects for a Document object that’s stored in a variable
named doc?
Ans.By using doc.paragraphs
!pip install python-docx
import docx
doc = docx.Document('abc.docx')
doc.paragraphs

8.What type of object has bold, underline, italic, strike, and outline variables?
Ans.A Run object

9.What is the difference between False, True, and None for the bold variable?
Ans.
True : the attribute is always enabled, no matter what other styles are applied to the run
        rue always makes the Run object bolded and False makes it always not bolded,
False : the attribute is always disabled
None : defaults to whatever the run’s style is set to

10.How do you create a Document object for a new Word document?
Ans.By Calling the docx.Document() function.

11.How do you add a paragraph with the text 'Hello, there' to a Document object stored in a
variable named doc?
Ans.
import docx
doc = docx.Document()
doc.add_paragraph('Hello there!')
doc.save('hellothere.docx')

12.What integers represent the levels of headings available in Word documents?
Ans.integer from 0 to 4
The arguments to add_heading() are a string of the heading text and an integer from 0 to 4.
The integer 0 makes the heading the Title style, which is used for the top of the document.
Integers 1 to 4 are for various heading levels, with 1 being the main heading and 4 the lowest subheading



'''